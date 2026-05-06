import streamlit as st

from io import BytesIO

try:
    from docx import Document
except ModuleNotFoundError:
    Document = None

try:
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except ModuleNotFoundError:
    getSampleStyleSheet = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None

try:
    from pypdf import PdfReader
except ModuleNotFoundError:
    PdfReader = None

from services.api_client import (
    analyze_cv,
    generate_cv,
    generate_cv_from_analysis,
    get_results,
    login_user,
    match_job,
    mock_interview,
    prepare_interview,
    register_user,
    save_result,
    search_jobs,
)

st.set_page_config(
    page_title="CareerBridge UK",
    page_icon="🇬🇧",
    layout="wide",
)


def extract_pdf_text(uploaded_file):
    if PdfReader is None:
        st.error("PDF upload support requires pypdf. Install frontend dependencies first.")
        return ""

    reader = PdfReader(uploaded_file)
    text = ""

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return text.strip()


def extract_docx_text(uploaded_file):
    if Document is None:
        st.error("DOCX upload support requires python-docx. Install frontend dependencies first.")
        return ""

    doc = Document(uploaded_file)
    text = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())

    return "\n".join(text)


def extract_cv_text(uploaded_file):
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return extract_pdf_text(uploaded_file)
    if file_name.endswith(".docx"):
        return extract_docx_text(uploaded_file)
    if file_name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")

    return ""


def call_backend(action, payload, result_key):
    try:
        response = action(payload)
    except RuntimeError as exc:
        st.error(str(exc))
        return None

    if result_key not in response:
        st.error(f"Unexpected backend response: {response}")
        return None

    return response[result_key]


def save_current_result(feature_type, input_text, result_text, target_role, location):
    try:
        response = save_result(
            {
                "feature_type": feature_type,
                "target_role": target_role,
                "location": location,
                "input_text": input_text,
                "result_text": result_text,
            }
        )
    except RuntimeError as exc:
        st.error(str(exc))
        return

    st.success(response.get("message", "Result saved successfully"))


def create_docx_download(text, title="CareerBridge UK Result"):
    if Document is None:
        return None

    doc = Document()
    doc.add_heading(title, level=1)

    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def create_pdf_download(text, title="CareerBridge UK CV"):
    if SimpleDocTemplate is None:
        return None

    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 12))

    for line in text.split("\n"):
        clean_line = line.strip()
        if clean_line:
            story.append(Paragraph(clean_line, styles["BodyText"]))
            story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)

    return buffer


def auth_screen():
    st.title("🇬🇧 CareerBridge UK")
    st.subheader("Login or create your account")

    auth_tab1, auth_tab2 = st.tabs(["Login", "Register"])

    with auth_tab1:
        email = st.text_input("Email", key="login_email")
        password = st.text_input("Password", type="password", key="login_password")

        if st.button("Login"):
            if not email.strip() or not password.strip():
                st.warning("Enter your email and password.")
                return

            response = login_user(
                {
                    "email": email.strip(),
                    "password": password.strip(),
                }
            )

            if "access_token" in response:
                st.session_state["logged_in"] = True
                st.session_state["token"] = response["access_token"]
                st.session_state["email"] = response["email"]
                st.success("Login successful")
                st.rerun()
            else:
                st.error(response.get("detail", "Login failed"))

    with auth_tab2:
        full_name = st.text_input("Full Name")
        reg_email = st.text_input("Email", key="register_email")
        reg_password = st.text_input(
            "Password",
            type="password",
            key="register_password",
        )

        if st.button("Create Account"):
            if not reg_email.strip() or not reg_password.strip():
                st.warning("Enter an email and password to create an account.")
                return

            response = register_user(
                {
                    "full_name": full_name.strip(),
                    "email": reg_email.strip(),
                    "password": reg_password.strip(),
                }
            )

            if "access_token" in response:
                st.session_state["logged_in"] = True
                st.session_state["token"] = response["access_token"]
                st.session_state["email"] = response["email"]
                st.success("Account created")
                st.rerun()
            else:
                st.error(response.get("detail", "Registration failed"))


if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

if not st.session_state["logged_in"]:
    auth_screen()
    st.stop()

st.markdown(
    """
<style>
.main {
    background-color: #f8fafc;
}

.block-container {
    padding-top: 2rem;
}

.hero-card {
    background: linear-gradient(135deg, #0f172a, #1e293b);
    padding: 32px;
    border-radius: 18px;
    color: white;
    margin-bottom: 24px;
}

.hero-title {
    font-size: 38px;
    font-weight: 800;
    margin-bottom: 8px;
}

.hero-subtitle {
    font-size: 18px;
    color: #cbd5e1;
}

.metric-card {
    background: white;
    padding: 20px;
    border-radius: 14px;
    border: 1px solid #e2e8f0;
    box-shadow: 0px 4px 12px rgba(15, 23, 42, 0.06);
}

.section-card {
    background: white;
    padding: 22px;
    border-radius: 16px;
    border: 1px solid #e2e8f0;
    margin-bottom: 16px;
}
</style>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<div class="hero-card">
    <div class="hero-title">🇬🇧 CareerBridge UK</div>
    <div class="hero-subtitle">
        AI Career Mentor for CV improvement, job matching, and interview preparation.
    </div>
</div>
""",
    unsafe_allow_html=True,
)

m1, m2, m3, m4 = st.columns(4)

with m1:
    st.markdown(
        '<div class="metric-card"><b>CV Analyzer</b><br>UK-standard feedback</div>',
        unsafe_allow_html=True,
    )

with m2:
    st.markdown(
        '<div class="metric-card"><b>Job Match</b><br>Role-fit scoring</div>',
        unsafe_allow_html=True,
    )

with m3:
    st.markdown(
        '<div class="metric-card"><b>Interview Prep</b><br>STAR answers</div>',
        unsafe_allow_html=True,
    )

with m4:
    st.markdown(
        '<div class="metric-card"><b>Mock Interview</b><br>AI interviewer</div>',
        unsafe_allow_html=True,
    )

with st.sidebar:
    st.header("Candidate Details")
    uploaded_cv = st.file_uploader("Upload your CV", type=["pdf", "docx", "txt"])
    manual_cv_text = st.text_area("Or paste your CV manually", height=200)
    target_role = st.text_input("Target role", "Warehouse Operative")
    location = st.text_input("Location", "Doncaster")
    experience_level = st.selectbox(
        "Experience level",
        ["Entry-level", "Some experience", "Experienced", "Career change"],
    )
    st.divider()
    if st.button("Logout", use_container_width=True):
        st.session_state["logged_in"] = False
        st.session_state.pop("token", None)
        st.session_state.pop("email", None)
        st.rerun()

cv_text = ""

if uploaded_cv:
    cv_text = extract_cv_text(uploaded_cv)
    if cv_text:
        st.success("CV uploaded successfully.")
        with st.expander("Preview extracted CV text"):
            st.text_area("Extracted CV", cv_text, height=250)
elif manual_cv_text:
    cv_text = manual_cv_text

base_payload = {
    "cv_text": cv_text,
    "target_role": target_role,
    "location": location,
    "experience_level": experience_level,
}

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "CV Tools",
    "Live Job Search",
    "Job Match",
    "Interview Preparation",
    "AI Mock Interview",
    "Saved Results"
])

with tab1:
    col1, col2 = st.columns(2)

    with col1:
        if st.button("Analyze CV", use_container_width=True):
            if not cv_text.strip():
                st.warning("Upload or paste your CV first.")
            else:
                result = call_backend(analyze_cv, base_payload, "analysis")
                if result:
                    st.session_state.cv_analysis_result = result

        if "cv_analysis_result" in st.session_state:
            st.subheader("CV Analysis")
            st.write(st.session_state.cv_analysis_result)
            if st.button("Generate Improved CV From This Analysis", use_container_width=True):
                if not cv_text.strip():
                    st.warning("Upload or paste your CV first.")
                else:
                    result = call_backend(
                        generate_cv_from_analysis,
                        {
                            **base_payload,
                            "cv_analysis": st.session_state.cv_analysis_result,
                        },
                        "improved_cv",
                    )
                    if result:
                        st.session_state.improved_cv_result = result

            if st.button("Save CV Analysis", use_container_width=True):
                save_current_result(
                    "CV Analysis",
                    cv_text,
                    st.session_state.cv_analysis_result,
                    target_role,
                    location,
                )

    with col2:
        if st.button("Generate UK CV", use_container_width=True):
            if not cv_text.strip():
                st.warning("Upload or paste your CV first.")
            else:
                result = call_backend(generate_cv, base_payload, "generated_cv")
                if result:
                    st.session_state.generated_cv_result = result

        if "generated_cv_result" in st.session_state:
            st.subheader("Generated UK CV")
            st.write(st.session_state.generated_cv_result)
            generated_cv_docx = create_docx_download(
                st.session_state.generated_cv_result,
                "Generated UK CV",
            )
            if generated_cv_docx:
                st.download_button(
                    label="Download CV as DOCX",
                    data=generated_cv_docx,
                    file_name="careerbridge_generated_cv.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            st.download_button(
                label="Download CV as TXT",
                data=st.session_state.generated_cv_result,
                file_name="careerbridge_generated_cv.txt",
                mime="text/plain",
                use_container_width=True,
            )
            if st.button("Save Generated CV", use_container_width=True):
                save_current_result(
                    "Generated CV",
                    cv_text,
                    st.session_state.generated_cv_result,
                    target_role,
                    location,
                )

        if "improved_cv_result" in st.session_state:
            st.subheader("Improved CV Based on Analysis")
            st.write(st.session_state.improved_cv_result)
            improved_cv_pdf = create_pdf_download(
                st.session_state.improved_cv_result,
                "Improved UK CV",
            )
            improved_cv_docx = create_docx_download(
                st.session_state.improved_cv_result,
                "Improved UK CV",
            )
            if improved_cv_pdf:
                st.download_button(
                    label="Download Improved CV as PDF",
                    data=improved_cv_pdf,
                    file_name="careerbridge_improved_cv.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            else:
                st.info("PDF download support requires reportlab.")

            if improved_cv_docx:
                st.download_button(
                    label="Download Improved CV as DOCX",
                    data=improved_cv_docx,
                    file_name="careerbridge_improved_cv.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            st.download_button(
                label="Download Improved CV as TXT",
                data=st.session_state.improved_cv_result,
                file_name="careerbridge_improved_cv.txt",
                mime="text/plain",
                use_container_width=True,
            )
            if st.button("Save Improved CV", use_container_width=True):
                save_current_result(
                    "Improved CV",
                    cv_text,
                    st.session_state.improved_cv_result,
                    target_role,
                    location,
                )

with tab2:
    st.header("Live UK Job Search")

    job_query = st.text_input(
        "Search job title or keyword",
        value=target_role,
        placeholder="e.g. warehouse operative, care assistant, support worker",
    )

    job_location = st.text_input(
        "Search location",
        value=location,
        placeholder="e.g. Doncaster, London, Manchester",
    )

    results_per_page = st.slider("Number of results", 5, 30, 10)

    if st.button("Search Live Jobs"):
        response = search_jobs(
            {
                "query": job_query,
                "location": job_location,
                "page": 1,
                "results_per_page": results_per_page,
            }
        )

        if "detail" in response:
            st.error(response["detail"])
        else:
            st.session_state.live_jobs = response["results"]
            st.session_state.live_jobs_count = response["count"]
            st.success(f"Found approximately {response['count']} jobs")

    for index, job in enumerate(st.session_state.get("live_jobs", [])):
        job_key = job.get("redirect_url") or f"{job.get('title')}_{index}"

        with st.expander(f"{job.get('title')} — {job.get('company')}"):
            st.write(f"**Location:** {job.get('location')}")
            st.write(
                f"**Salary:** {job.get('salary_min')} - {job.get('salary_max')}"
            )
            st.write(job.get("description"))

            col_a, col_b, col_c = st.columns(3)

            with col_a:
                if st.button(
                    "Check CV Match",
                    key=f"match_{job_key}",
                    use_container_width=True,
                ):
                    if not cv_text.strip():
                        st.warning("Upload or paste your CV first.")
                    else:
                        result = call_backend(
                            match_job,
                            {
                                "cv_text": cv_text,
                                "target_role": job.get("title"),
                                "job_description": job.get("description"),
                                "location": job.get("location"),
                                "experience_level": experience_level,
                            },
                            "match_result",
                        )
                        if result:
                            st.session_state[f"live_job_match_{job_key}"] = result

            with col_b:
                if st.button(
                    "Generate Tailored CV",
                    key=f"tailored_cv_{job_key}",
                    use_container_width=True,
                ):
                    if not cv_text.strip():
                        st.warning("Upload or paste your CV first.")
                    else:
                        result = call_backend(
                            generate_cv_from_analysis,
                            {
                                "cv_text": cv_text,
                                "cv_analysis": f"""
                                Tailor this CV specifically for this job.

                                Job title: {job.get("title")}
                                Company: {job.get("company")}
                                Location: {job.get("location")}
                                Job description: {job.get("description")}
                                """,
                                "target_role": job.get("title"),
                                "location": job.get("location"),
                                "experience_level": experience_level,
                            },
                            "improved_cv",
                        )
                        if result:
                            st.session_state[f"live_job_tailored_cv_{job_key}"] = result

            with col_c:
                if st.button(
                    "Prepare Interview",
                    key=f"interview_{job_key}",
                    use_container_width=True,
                ):
                    if not cv_text.strip():
                        st.warning("Upload or paste your CV first.")
                    else:
                        result = call_backend(
                            prepare_interview,
                            {
                                "cv_text": cv_text,
                                "target_role": job.get("title"),
                                "job_description": job.get("description"),
                                "location": job.get("location"),
                                "experience_level": experience_level,
                                "interview_style": "Formal",
                            },
                            "preparation",
                        )
                        if result:
                            st.session_state[f"live_job_interview_{job_key}"] = result

            if f"live_job_match_{job_key}" in st.session_state:
                st.subheader("AI Job Match Result")
                st.write(st.session_state[f"live_job_match_{job_key}"])

            if f"live_job_tailored_cv_{job_key}" in st.session_state:
                tailored_cv = st.session_state[f"live_job_tailored_cv_{job_key}"]
                st.subheader("Tailored CV for This Job")
                st.write(tailored_cv)

                tailored_pdf = create_pdf_download(tailored_cv, "Tailored UK CV")
                tailored_docx = create_docx_download(tailored_cv, "Tailored UK CV")

                if tailored_pdf:
                    st.download_button(
                        "Download Tailored CV as PDF",
                        data=tailored_pdf,
                        file_name="careerbridge_tailored_cv.pdf",
                        mime="application/pdf",
                        key=f"download_pdf_{job_key}",
                        use_container_width=True,
                    )
                else:
                    st.info("PDF download support requires reportlab.")

                if tailored_docx:
                    st.download_button(
                        "Download Tailored CV as DOCX",
                        data=tailored_docx,
                        file_name="careerbridge_tailored_cv.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_docx_{job_key}",
                        use_container_width=True,
                    )

            if f"live_job_interview_{job_key}" in st.session_state:
                st.subheader("Interview Preparation for This Job")
                st.write(st.session_state[f"live_job_interview_{job_key}"])

            if job.get("redirect_url"):
                st.link_button("View / Apply", job.get("redirect_url"))

with tab3:
    job_description = st.text_area("Paste the job description", height=180)

    if st.button("Match Job", use_container_width=True):
        if not cv_text.strip() or not job_description.strip():
            st.warning("Upload or paste your CV and job description first.")
        else:
            result = call_backend(
                match_job,
                {
                    **base_payload,
                    "job_description": job_description,
                },
                "match_result",
            )
            if result:
                st.session_state.job_match_result = result
                st.session_state.job_match_input = job_description

    if "job_match_result" in st.session_state:
        st.subheader("Job Match")
        st.write(st.session_state.job_match_result)
        if st.button("Save Job Match Result", use_container_width=True):
            save_current_result(
                "Job Match",
                st.session_state.get("job_match_input", job_description),
                st.session_state.job_match_result,
                target_role,
                location,
            )

with tab4:
    interview_style = st.selectbox(
        "Interview style",
        [
            "Formal",
            "Friendly",
            "Strict",
            "NHS style",
            "Warehouse recruiter style",
            "Care sector style",
            "AI recruiter style",
        ],
    )
    interview_job_description = st.text_area(
        "Paste job description optional",
        height=160,
        key="interview_job_description",
    )

    if st.button("Prepare Interview Answers", use_container_width=True):
        if not cv_text.strip():
            st.warning("Upload or paste your CV first.")
        else:
            result = call_backend(
                prepare_interview,
                {
                    **base_payload,
                    "job_description": interview_job_description,
                    "interview_style": interview_style,
                },
                "preparation",
            )
            if result:
                st.session_state.interview_prep_result = result
                st.session_state.interview_prep_input = interview_job_description

    if "interview_prep_result" in st.session_state:
        st.subheader("Interview Preparation")
        st.write(st.session_state.interview_prep_result)
        interview_docx = create_docx_download(
            st.session_state.interview_prep_result,
            "Interview Preparation",
        )
        if interview_docx:
            st.download_button(
                label="Download Interview Prep as DOCX",
                data=interview_docx,
                file_name="careerbridge_interview_prep.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        st.download_button(
            label="Download Interview Prep as TXT",
            data=st.session_state.interview_prep_result,
            file_name="careerbridge_interview_prep.txt",
            mime="text/plain",
            use_container_width=True,
        )
        if st.button("Save Interview Preparation", use_container_width=True):
            save_current_result(
                "Interview Preparation",
                st.session_state.get("interview_prep_input", interview_job_description),
                st.session_state.interview_prep_result,
                target_role,
                location,
            )

with tab5:
    mock_interview_style = st.selectbox(
        "Interview style",
        [
            "Formal",
            "Friendly",
            "Strict",
            "NHS style",
            "Warehouse recruiter style",
            "Care sector style",
        ],
        key="mock_interview_style",
    )

    if "mock_messages" not in st.session_state:
        st.session_state.mock_messages = [
            {
                "role": "assistant",
                "content": "Hello, thank you for attending this interview. Can you tell me about yourself?",
            }
        ]

    for message in st.session_state.mock_messages:
        role = "assistant" if message["role"] == "assistant" else "user"
        with st.chat_message(role):
            st.write(message["content"])

    user_answer = st.chat_input("Type your interview answer...")

    if user_answer:
        if not cv_text.strip():
            st.warning("Upload or paste your CV first.")
        else:
            st.session_state.mock_messages.append(
                {
                    "role": "user",
                    "content": user_answer,
                }
            )

            try:
                response = mock_interview(
                    {
                        **base_payload,
                        "interview_style": mock_interview_style,
                        "messages": st.session_state.mock_messages,
                    }
                )
            except RuntimeError as exc:
                st.error(str(exc))
            else:
                ai_reply = response.get("reply")
                if ai_reply:
                    st.session_state.mock_messages.append(
                        {
                            "role": "assistant",
                            "content": ai_reply,
                        }
                    )
                    st.rerun()
                else:
                    st.error(f"Unexpected backend response: {response}")

    if st.button("Reset Mock Interview", use_container_width=True):
        st.session_state.mock_messages = [
            {
                "role": "assistant",
                "content": "Hello, thank you for attending this interview. Can you tell me about yourself?",
            }
        ]
        st.rerun()

with tab6:
    if st.button("Refresh Saved Results", use_container_width=True):
        try:
            st.session_state.saved_results = get_results()
        except RuntimeError as exc:
            st.error(str(exc))

    for result in st.session_state.get("saved_results", []):
        with st.expander(
            f"{result.get('feature_type', 'Result')} - {result.get('target_role') or 'No role'}"
        ):
            st.caption(result.get("created_at", ""))
            if result.get("location"):
                st.write(f"Location: {result['location']}")
            st.write(result.get("result_text", ""))
